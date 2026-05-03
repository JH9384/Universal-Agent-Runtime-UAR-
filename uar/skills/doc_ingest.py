import os
from pathlib import Path
from typing import Any, Generator
import logging

from uar.core.registry import register_skill
from uar.core.exceptions import PathSecurityError
from uar.core.validation import validate_path_security

logger = logging.getLogger(__name__)

# Resolve allowed root from environment or use current working directory
# In production, set PROJECT_ROOT env var to ensure consistent path resolution
_allowed_root_env = os.getenv("PROJECT_ROOT") or os.getenv("RUNS_DIR")
ALLOWED_ROOT = Path(_allowed_root_env).resolve() if _allowed_root_env else Path.cwd()
ALLOWED_EXTENSIONS = {
    # docs / text
    ".txt", ".md", ".rst", ".markdown", ".log",
    ".adoc", ".asciidoc", ".org", ".tex", ".ltx", ".sty", ".cls",
    ".bib", ".ris",
    # config
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env", ".conf",
    # python / scripting
    ".py", ".pyi", ".pyx", ".sh", ".bash", ".zsh", ".rb", ".pl", ".lua",
    # web
    ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs",
    ".html", ".htm", ".xml", ".css", ".scss", ".sass", ".less",
    ".vue", ".svelte",
    # systems / compiled
    ".c", ".h", ".cpp", ".hpp", ".cc", ".cxx",
    ".rs", ".go", ".java", ".kt", ".swift", ".m", ".mm",
    ".cs", ".php", ".scala", ".clj", ".ex", ".exs",
    # data / sql / serialization
    ".sql", ".graphql", ".proto",
    ".csv", ".tsv", ".dat", ".data", ".jsonl", ".ndjson", ".geojson",
    # build
    ".dockerfile", ".makefile", ".mk", ".gradle", ".cmake",
    # science / math / stats
    ".ipynb",                          # Jupyter (JSON, read raw)
    ".r", ".rmd", ".rmarkdown", ".qmd", # R, R-Markdown, Quarto
    ".jl",                              # Julia
    ".lean", ".v", ".thy",              # proof assistants
    ".m",                               # MATLAB / Octave (also Obj-C above; both OK)
    ".sage",                            # SageMath
    ".gp",                              # PARI/GP
    ".mac", ".max",                     # Maxima
    ".asy",                             # Asymptote
    ".mps", ".lp",                      # LP/MPS optimization
    ".smt", ".smt2",                    # SMT-LIB
    ".pddl",                            # planning
    ".g4",                              # ANTLR grammars
    # binary docs / data (handled by special extractors)
    ".pdf",                             # pypdf
    ".docx",                            # python-docx
    ".xlsx", ".xlsm",                   # openpyxl
    ".parquet", ".feather",             # pandas + pyarrow
}

# Still NOT supported (would need extra deps):
# .mat .nb (Mathematica) .fits .h5 .hdf5 .nc .npz .npy
# .doc (legacy Word) .xls (legacy Excel) .pptx (python-pptx)
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILES = 1000
MAX_TOTAL_SIZE = 100 * 1024 * 1024  # 100MB total limit to prevent memory exhaustion


def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF using pypdf if available. Returns '' on failure."""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            raise RuntimeError("pypdf not installed (pip install pypdf)")
    reader = PdfReader(str(file_path))
    out = []
    for i, page in enumerate(reader.pages):
        try:
            out.append(page.extract_text() or "")
        except Exception as e:
            out.append(f"[page {i+1} extraction error: {e}]")
    return "\n\n".join(out)


def _extract_ipynb(file_path: Path) -> str:
    """Strip a Jupyter notebook to markdown + fenced code via nbformat (with fallback)."""
    try:
        import nbformat  # type: ignore
        nb = nbformat.read(str(file_path), as_version=4)
        cells = nb.get("cells", [])
        lang = (nb.get("metadata", {}).get("kernelspec", {}).get("language") or "python").lower()
    except Exception:
        import json
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                nb = json.load(f)
            cells = nb.get("cells", [])
            lang = "python"
        except Exception as e:
            return f"[ipynb parse failed: {e}]"
    parts = []
    for cell in cells:
        ctype = cell.get("cell_type", "")
        src = cell.get("source", "")
        if isinstance(src, list):
            src = "".join(src)
        if not src.strip():
            continue
        if ctype == "markdown":
            parts.append(src)
        elif ctype == "code":
            parts.append(f"```{lang}\n{src}\n```")
        else:
            parts.append(src)
    return "\n\n".join(parts)


def _extract_docx(file_path: Path) -> str:
    """Extract paragraphs + tables from a .docx via python-docx."""
    from docx import Document  # type: ignore
    doc = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(file_path: Path) -> str:
    """Extract all sheets from a .xlsx as TSV-like text via openpyxl."""
    from openpyxl import load_workbook  # type: ignore
    wb = load_workbook(str(file_path), data_only=True, read_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"--- SHEET: {sheet} ---")
        rows_added = 0
        for row in ws.iter_rows(values_only=True):
            if all(c is None for c in row):
                continue
            parts.append("\t".join("" if c is None else str(c) for c in row))
            rows_added += 1
            if rows_added >= 5000:
                parts.append("[...truncated at 5000 rows...]")
                break
    wb.close()
    return "\n".join(parts)


def _extract_dataframe(file_path: Path, kind: str) -> str:
    """Extract a parquet/feather file as CSV-like text via pandas."""
    import pandas as pd  # type: ignore
    if kind == "parquet":
        df = pd.read_parquet(str(file_path))
    elif kind == "feather":
        df = pd.read_feather(str(file_path))
    else:
        raise ValueError(f"Unsupported dataframe kind: {kind}")
    head = df.head(2000)
    info = f"# {kind} · shape={df.shape} · cols={list(df.columns)}\n"
    return info + head.to_csv(index=False)


def _read_file_safely(file_path: Path, allowed_root: Path) -> dict[str, Any]:
    """Read a single file with proper resource management and security checks.

    Uses context manager to ensure file handle is closed even on errors.
    Special-cases PDF and Jupyter notebooks.
    """
    try:
        # Validate security before reading
        validate_path_security(file_path, allowed_root)
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE:
            return {
                "path": str(file_path.relative_to(allowed_root)),
                "text": "",
                "size": 0,
                "error": f"File too large: {file_size} bytes (max {MAX_FILE_SIZE})"
            }

        suffix = file_path.suffix.lower()

        # Binary special-cased extractors
        binary_extractors = {
            ".pdf":     ("pdf",     _extract_pdf),
            ".docx":    ("docx",    _extract_docx),
            ".xlsx":    ("xlsx",    _extract_xlsx),
            ".xlsm":    ("xlsx",    _extract_xlsx),
            ".ipynb":   ("ipynb",   _extract_ipynb),
            ".parquet": ("parquet", lambda p: _extract_dataframe(p, "parquet")),
            ".feather": ("feather", lambda p: _extract_dataframe(p, "feather")),
        }
        if suffix in binary_extractors:
            kind, fn = binary_extractors[suffix]
            try:
                content = fn(file_path)
                return {
                    "path": str(file_path.relative_to(allowed_root)),
                    "text": content,
                    "size": len(content),
                    "type": kind,
                }
            except Exception as e:
                return {
                    "path": str(file_path.relative_to(allowed_root)),
                    "text": "",
                    "size": 0,
                    "error": f"{kind} extraction failed: {e}",
                }

        # Plain text path
        with open(file_path, 'r', encoding='utf-8', errors='strict') as f:
            content = f.read()
        return {
            "path": str(file_path.relative_to(allowed_root)),
            "text": content,
            "size": len(content),
            "type": suffix.lstrip("."),
        }
        
    except UnicodeDecodeError:
        return {
            "path": str(file_path.relative_to(allowed_root)) if file_path.is_relative_to(allowed_root) else str(file_path),
            "text": "",
            "size": 0,
            "error": "File encoding error (not valid UTF-8)"
        }
    except PermissionError:
        return {
            "path": str(file_path.relative_to(allowed_root)) if file_path.is_relative_to(allowed_root) else str(file_path),
            "text": "",
            "size": 0,
            "error": "Permission denied"
        }
    except Exception as e:
        return {
            "path": str(file_path.relative_to(allowed_root)) if file_path.is_relative_to(allowed_root) else str(file_path),
            "text": "",
            "size": 0,
            "error": f"Read error: {str(e)}"
        }


def _yield_documents(path: Path, allowed_root: Path) -> Generator[dict[str, Any], None, None]:
    """Generator to yield documents one at a time to avoid loading all into memory."""
    file_count = 0
    total_size = 0
    
    if path.is_file():
        if path.suffix.lower() in ALLOWED_EXTENSIONS:
            doc = _read_file_safely(path, allowed_root)
            # Always yield the doc, even if it has an error (so caller knows what happened)
            yield doc
        else:
            yield {
                "path": str(path.relative_to(allowed_root)) if path.is_relative_to(allowed_root) else str(path),
                "text": "",
                "size": 0,
                "error": f"Unsupported file type: {path.suffix}"
            }
    elif path.is_dir():
        # Stream entries lazily; do not materialize the full tree via sorted().
        for entry in path.rglob("*"):
            if file_count >= MAX_FILES:
                yield {
                    "path": "LIMIT_EXCEEDED", 
                    "text": f"Stopped at {MAX_FILES} files",
                    "size": 0,
                    "warning": "Maximum file count reached"
                }
                return
            
            if total_size >= MAX_TOTAL_SIZE:
                yield {
                    "path": "SIZE_LIMIT_EXCEEDED",
                    "text": f"Stopped at {total_size} bytes",
                    "size": 0,
                    "warning": "Maximum total size reached"
                }
                return
                
            if entry.is_file() and entry.suffix.lower() in ALLOWED_EXTENSIONS:
                try:
                    # Quick size check before reading
                    entry_size = entry.stat().st_size
                    if entry_size > MAX_FILE_SIZE:
                        yield {
                            "path": str(entry.relative_to(allowed_root)) if entry.is_relative_to(allowed_root) else str(entry),
                            "text": "",
                            "size": 0,
                            "error": f"File too large: {entry_size} bytes"
                        }
                        file_count += 1
                        continue
                    
                    doc = _read_file_safely(entry, allowed_root)
                    file_count += 1
                    
                    if "error" not in doc or doc["error"] == "":
                        total_size += doc.get("size", 0)
                    
                    yield doc
                    
                except OSError as e:
                    # Handle race conditions where file is deleted between check and read
                    logger.warning(f"File access error for {entry}: {e}")
                    continue
    else:
        yield {
            "path": str(path),
            "text": "",
            "size": 0,
            "error": f"Input path not found: {path}"
        }


@register_skill("doc_ingest")
def doc_ingest(ctx):
    """Ingest documents from a specified path with security validation.
    
    Features:
    - Proper resource management with context managers
    - Memory-efficient streaming via generators
    - Total size limits to prevent memory exhaustion
    - Comprehensive error handling for individual files
    - Security validation at every access point
    """
    input_path = ctx.goal.metadata.get("input_path")
    if not input_path:
        return {"documents": [], "warning": "No input_path provided"}

    try:
        path = Path(input_path).resolve()
        
        # Validate path security before any access
        validate_path_security(path, ALLOWED_ROOT)
        
        # Use generator and convert to list with size limits enforced
        documents = []
        total_size = 0
        doc_count = 0
        
        for doc in _yield_documents(path, ALLOWED_ROOT):
            # If we've already hit the limit, this doc should be LIMIT_EXCEEDED
            if doc_count >= MAX_FILES:
                documents.append(doc)  # This is the LIMIT_EXCEEDED marker
                break
            
            documents.append(doc)
            doc_count += 1
            
            # Track size only for successfully read documents
            if "error" not in doc or doc["error"] == "":
                total_size += doc.get("size", 0)
            
        return {
            "documents": documents,
            "document_count": len([d for d in documents if "error" not in d or d.get("error") == ""]),
            "paths": [doc["path"] for doc in documents if "error" not in doc or doc.get("error") == ""],
            "total_size": total_size,
            "errors": [doc["error"] for doc in documents if "error" in doc and doc["error"]],
        }
        
    except PathSecurityError as e:
        logger.error(f"Path security error: {e}")
        return {"documents": [], "error": str(e)}
    except Exception as e:
        logger.error(f"Unexpected error in doc_ingest: {e}", exc_info=True)
        return {"documents": [], "error": f"Unexpected error: {str(e)}"}
